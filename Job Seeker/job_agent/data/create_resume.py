#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

# Header - Name
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name.add_run('Piotr Kaminski')
name_run.font.size = Pt(14)
name_run.font.bold = True

# Contact
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.add_run('+48 510 934 702 | piotrkaminski1@protonmail.com | linkedin.com/in/piotr-kaminski | Krakow, Poland').font.size = Pt(9)

# Summary
summary_h = doc.add_paragraph()
summary_h.add_run('PROFESSIONAL SUMMARY').font.bold = True
summary_h.runs[0].font.size = Pt(11)

summary = doc.add_paragraph('Technical Solutions Engineer with 10+ years driving API integration, SaaS product adoption, and customer success. Proven ability to troubleshoot complex technical issues, design customer onboarding processes, and deliver measurable business impact through cross-functional collaboration and data-driven problem solving.')
summary.paragraph_format.space_after = Pt(6)

# Experience
exp_h = doc.add_paragraph()
exp_h.add_run('EXPERIENCE').font.bold = True
exp_h.runs[0].font.size = Pt(11)

# Job 1: Lead Solutions Engineer
j1_title = doc.add_paragraph()
j1_title.add_run('Lead Solutions Engineer').font.bold = True
j1_title.add_run(' | Luigi\'s Box | Jan 2025–Present')

bullets_j1 = [
    'Achieved zero churn rate from integration implementations through comprehensive customer enablement and technical onboarding',
    'Led customer onboarding and integration initiatives for key accounts, ensuring 100% successful implementations',
    'Redesigned technical onboarding processes to improve customer success rates and reduce implementation friction',
    'Managed cross-functional collaboration between engineering and customer success teams; developed technical documentation for product enablement',
    'Conducted SQL data analysis to identify customer usage patterns and upsell opportunities'
]

for bullet in bullets_j1:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

# Job 2: Technical Account Manager
j2_title = doc.add_paragraph()
j2_title.add_run('Technical Account Manager').font.bold = True
j2_title.add_run(' | Flexiroam | Jan 2023–Dec 2024')

bullets_j2 = [
    'Delivered 100% implementation success rate across all B2B migrations and customer integrations',
    'Grew assigned accounts by 25% year-over-year through technical enablement and customer support excellence',
    'Developed and maintained REST API integrations for enterprise clients; provided API troubleshooting and debugging support',
    'Led successful B2B platform migration affecting 50+ accounts; coordinated cross-functional team alignment',
    'Leveraged SQL and data analysis to drive customer adoption insights and upsell strategy'
]

for bullet in bullets_j2:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

# Job 3: Technical Support Engineer
j3_title = doc.add_paragraph()
j3_title.add_run('Technical Support Engineer').font.bold = True
j3_title.add_run(' | Rapid | Jan 2020–Dec 2022')

bullets_j3 = [
    'Managed Tier 2 technical support with 89% first-contact satisfaction rate; expertise in API troubleshooting and debugging complex integrations',
    'Developed and documented technical solutions for developer platform integration issues; created support documentation and runbooks',
    'Trained and mentored 7 engineers on API integration troubleshooting, technical support best practices, and customer communication',
    'Optimized support workflows and debugging processes, improving operational efficiency by 10-15% and reducing mean-time-to-resolution by 20%',
    'Developed fraud detection mechanisms protecting $300K+ in platform transactions through SQL query analysis'
]

for bullet in bullets_j3:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

# Skills
skills_h = doc.add_paragraph()
skills_h.add_run('TECHNICAL SKILLS & TOOLS').font.bold = True
skills_h.runs[0].font.size = Pt(11)

sk1 = doc.add_paragraph()
sk1.add_run('APIs & Integration: ').font.bold = True
sk1.add_run('REST APIs, API troubleshooting, API integration, Webhooks, Integration support')

sk2 = doc.add_paragraph()
sk2.add_run('Tools & Platforms: ').font.bold = True
sk2.add_run('Postman, SQL, JSON, CloudWatch, Grafana, JIRA, Confluence, VS Code')

sk3 = doc.add_paragraph()
sk3.add_run('Languages: ').font.bold = True
sk3.add_run('English (C1/C2 professional fluency), Polish (native)')

sk4 = doc.add_paragraph()
sk4.add_run('Core Competencies: ').font.bold = True
sk4.add_run('Technical troubleshooting, Customer onboarding, SaaS product knowledge, Cross-functional collaboration, Technical documentation, Problem solving, Process optimization, Team mentoring')

doc.save('resumes/cv_linkedin_blz_001.docx')
print("✓ Resume created")
