#!/usr/bin/env python3
import sys
sys.path.insert(0, '/Users/bill/Work/codeworkspace/.claude/worktrees/interesting-swanson/job_agent/data')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# CREATE TAILORED RESUME FOR CONSTRUCTOR
# ============================================================================
doc_resume = Document()

for section in doc_resume.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Header with name
name = doc_resume.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name.add_run('PIOTR KAMINSKI')
name_run.font.size = Pt(16)
name_run.font.bold = True

# Contact information
contact = doc_resume.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.add_run('Krakow, Poland | +48 510 934 702 | piotrkaminski1@protonmail.com | linkedin.com/in/piotr-kaminski')
contact.paragraph_format.space_after = Pt(6)

# Professional Summary
summary_h = doc_resume.add_paragraph()
summary_run = summary_h.add_run('PROFESSIONAL SUMMARY')
summary_run.font.bold = True
summary_run.font.size = Pt(11)

summary = doc_resume.add_paragraph(
    'Technical Integration Specialist with 10+ years driving SaaS platform adoption, enterprise API integration, and customer onboarding. '
    'Proven track record delivering zero churn from integration implementations, 100% B2B migration success rates, and 25% year-over-year account growth. '
    'Deep expertise in AI-powered ecommerce search platforms, REST API documentation, and complex enterprise onboarding at scale.'
)
summary.paragraph_format.space_after = Pt(12)

# Experience section
exp_h = doc_resume.add_paragraph()
exp_run = exp_h.add_run('EXPERIENCE')
exp_run.font.bold = True
exp_run.font.size = Pt(11)

# Job 1: Lead Solutions Engineer
j1_title = doc_resume.add_paragraph()
j1_title_run = j1_title.add_run('Lead Solutions Engineer, Luigi\'s Box')
j1_title_run.font.bold = True
j1_title.add_run(' | Remote | Jan 2025–Present')

for bullet in [
    'Redesigned SaaS customer onboarding processes, improving API integration efficiency and enterprise client success rates',
    'Achieved zero churn rate from API integration implementations—100% customer retention on all integration projects',
    'Led API Integration and Customer Onboarding initiatives for key enterprise accounts across ecommerce search platforms',
    'Managed Tier 1 and Tier 2 technical delivery, troubleshooting, and customer support for SaaS platform adoption',
    'Leveraged AI-powered ecommerce search and product data configuration expertise to accelerate customer time-to-value'
]:
    p = doc_resume.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)

# Job 2: Technical Account Manager
j2_title = doc_resume.add_paragraph()
j2_title_run = j2_title.add_run('Technical Account Manager, Flexiroam')
j2_title_run.font.bold = True
j2_title.add_run(' | Remote | Jan 2023–Dec 2024')

for bullet in [
    'Led successful B2B platform migration affecting 50+ enterprise accounts with 100% implementation success rate',
    'Developed and maintained REST API integrations for enterprise clients, enabling complex B2B technical workflows',
    'Mastered enterprise customer onboarding at scale, delivering comprehensive API documentation and product data configuration',
    'Grew assigned accounts by 25% year-over-year through technical depth, customer advocacy, and integration mastery'
]:
    p = doc_resume.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)

# Job 3: Technical Support Engineer
j3_title = doc_resume.add_paragraph()
j3_title_run = j3_title.add_run('Technical Support Engineer, Rapid')
j3_title_run.font.bold = True
j3_title.add_run(' | Remote | Jan 2020–Dec 2022')

for bullet in [
    'Managed Tier 2 technical support with 89% first-contact satisfaction rate and expert troubleshooting capabilities',
    'Optimized technical support workflows and problem-solving processes, improving operational efficiency by 10–15%',
    'Reduced mean-time-to-resolution by 20% through systematic technical troubleshooting and process improvements',
    'Developed fraud detection mechanisms protecting $300K+ in transactions, demonstrating technical depth in complex systems',
    'Trained and mentored 7 engineers on technical fundamentals and customer communication best practices'
]:
    p = doc_resume.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)

# Technical Skills section
skills_h = doc_resume.add_paragraph()
skills_run = skills_h.add_run('TECHNICAL SKILLS')
skills_run.font.bold = True
skills_run.font.size = Pt(11)

skills_text = doc_resume.add_paragraph(
    'REST APIs • API Integration • JSON • Python • SQL • Postman • Browser Developer Tools • CloudWatch • Grafana • JIRA • Confluence • E-commerce Platforms • Product Data Configuration • API Documentation • SaaS Platforms'
)
skills_text.paragraph_format.space_after = Pt(12)

# Languages section
lang_h = doc_resume.add_paragraph()
lang_run = lang_h.add_run('LANGUAGES')
lang_run.font.bold = True
lang_run.font.size = Pt(11)

lang_text = doc_resume.add_paragraph(
    'English (C1/C2 – Professional Fluency) | Polish (Native)'
)

doc_resume.save('../resumes/cv_linkedin_const_001.docx')
print("✓ Resume created: resumes/cv_linkedin_const_001.docx")

# ============================================================================
# CREATE TAILORED COVER LETTER FOR CONSTRUCTOR
# ============================================================================
doc_letter = Document()

for section in doc_letter.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Date
date_para = doc_letter.add_paragraph('April 14, 2026')
date_para.paragraph_format.space_after = Pt(12)

# Recipient info
recipient = doc_letter.add_paragraph('Hiring Team\nConstructor\n')
recipient.paragraph_format.space_after = Pt(12)

# Salutation
salutation = doc_letter.add_paragraph('Dear Hiring Manager,')
salutation.paragraph_format.space_after = Pt(12)

# Opening paragraph
opening = doc_letter.add_paragraph(
    'I\'m writing to express my interest in the Technical Integration Specialist role at Constructor. '
    'I\'m currently at Luigi\'s Box working on AI-powered ecommerce search—exactly the problem Constructor solves. '
    'This domain expertise would enable me to onboard your customers with immediate product knowledge and confidence, making me effective from day one.'
)
opening.paragraph_format.space_after = Pt(12)
opening.paragraph_format.line_spacing = 1.15

# Body paragraph 1
body1 = doc_letter.add_paragraph(
    'My integration and enterprise onboarding track record is exceptionally strong. At Luigi\'s Box, I\'ve achieved zero churn rate '
    'from API integration implementations—a metric that reflects obsessive attention to customer success in every technical interaction. '
    'At Flexiroam, I led a successful B2B platform migration affecting 50+ enterprise accounts with a 100% implementation success rate. '
    'This isn\'t just support; it\'s technical mastery applied to complex, high-stakes integrations.'
)
body1.paragraph_format.space_after = Pt(12)
body1.paragraph_format.line_spacing = 1.15

# Body paragraph 2
body2 = doc_letter.add_paragraph(
    'Throughout my 10+ years in SaaS, I\'ve learned that the best integration specialist combines technical depth with customer advocacy. '
    'I don\'t just troubleshoot API issues—I understand the business context driving integration needs and ensure clients extract real value '
    'from their implementation. This approach has driven 25% year-over-year account growth and consistently delivered outcomes that exceed typical specialist-level expectations.'
)
body2.paragraph_format.space_after = Pt(12)
body2.paragraph_format.line_spacing = 1.15

# Closing paragraph
closing = doc_letter.add_paragraph(
    'I\'m drawn to Constructor because your growth-stage position in the ecommerce search market aligns perfectly with my domain expertise '
    'and proven ability to accelerate customer success. I\'d welcome the opportunity to discuss how my integration mastery, API documentation expertise, '
    'and zero-churn track record can help Constructor build lasting customer relationships and scale your onboarding operation. '
    'Thank you for considering my application.'
)
closing.paragraph_format.space_after = Pt(12)
closing.paragraph_format.line_spacing = 1.15

# Sign-off
sign_off = doc_letter.add_paragraph('Best regards,')
sign_off.paragraph_format.space_after = Pt(24)

# Name
name_sig = doc_letter.add_paragraph('Piotr Kaminski')

doc_letter.save('../cover_letters/letter_linkedin_const_001.docx')
print("✓ Cover letter created: cover_letters/letter_linkedin_const_001.docx")

print("\n" + "="*70)
print("CONSTRUCTOR APPLICATION MATERIALS GENERATED SUCCESSFULLY")
print("="*70)
