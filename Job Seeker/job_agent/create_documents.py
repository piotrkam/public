#!/usr/bin/env python3
"""
Create tailored resume and cover letter DOCX files using python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_resume():
    """Create the tailored resume DOCX file"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Header with name
    name = doc.add_paragraph()
    name_run = name.add_run("PIOTR KAMINSKI")
    name_run.font.size = Pt(16)
    name_run.font.bold = True
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact information
    contact = doc.add_paragraph()
    contact_text = "Krakow, Poland | +48 510 934 702 | piotrkaminski1@protonmail.com | linkedin.com/in/piotr-kaminski"
    contact.add_run(contact_text)
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_format = contact.paragraph_format
    contact_format.space_after = Pt(6)

    # Professional Summary
    summary_heading = doc.add_paragraph()
    summary_run = summary_heading.add_run("PROFESSIONAL SUMMARY")
    summary_run.font.bold = True
    summary_run.font.size = Pt(11)

    summary_text = doc.add_paragraph(
        "Technical Integration Specialist with 10+ years driving SaaS platform adoption, enterprise API integration, and customer onboarding. "
        "Proven track record delivering zero churn from integration implementations, 100% B2B migration success rates, and 25% year-over-year account growth. "
        "Deep expertise in AI-powered ecommerce search platforms, REST API documentation, and complex enterprise onboarding at scale."
    )
    summary_text.paragraph_format.space_after = Pt(12)

    # Experience section
    exp_heading = doc.add_paragraph()
    exp_run = exp_heading.add_run("EXPERIENCE")
    exp_run.font.bold = True
    exp_run.font.size = Pt(11)

    # Job 1: Lead Solutions Engineer
    job1 = doc.add_paragraph()
    job1_title = job1.add_run("Lead Solutions Engineer, Luigi's Box")
    job1_title.font.bold = True
    job1.add_run(" | Remote | Jan 2025–Present")

    job1_bullets = [
        "Redesigned SaaS customer onboarding processes, improving API integration efficiency and enterprise client success rates",
        "Achieved zero churn rate from API integration implementations—100% customer retention on all integration projects",
        "Led API Integration and Customer Onboarding initiatives for key enterprise accounts across ecommerce search platforms",
        "Managed Tier 1 and Tier 2 technical delivery, troubleshooting, and customer support for SaaS platform adoption",
        "Leveraged AI-powered ecommerce search and product data configuration expertise to accelerate customer time-to-value"
    ]

    for bullet in job1_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)

    # Job 2: Technical Account Manager
    job2 = doc.add_paragraph()
    job2_title = job2.add_run("Technical Account Manager, Flexiroam")
    job2_title.font.bold = True
    job2.add_run(" | Remote | Jan 2023–Dec 2024")

    job2_bullets = [
        "Led successful B2B platform migration affecting 50+ enterprise accounts with 100% implementation success rate",
        "Developed and maintained REST API integrations for enterprise clients, enabling complex B2B technical workflows",
        "Mastered enterprise customer onboarding at scale, delivering comprehensive API documentation and product data configuration",
        "Grew assigned accounts by 25% year-over-year through technical depth, customer advocacy, and integration mastery"
    ]

    for bullet in job2_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)

    # Job 3: Technical Support Engineer
    job3 = doc.add_paragraph()
    job3_title = job3.add_run("Technical Support Engineer, Rapid")
    job3_title.font.bold = True
    job3.add_run(" | Remote | Jan 2020–Dec 2022")

    job3_bullets = [
        "Managed Tier 2 technical support with 89% first-contact satisfaction rate and expert troubleshooting capabilities",
        "Optimized technical support workflows and problem-solving processes, improving operational efficiency by 10–15%",
        "Reduced mean-time-to-resolution by 20% through systematic technical troubleshooting and process improvements",
        "Developed fraud detection mechanisms protecting $300K+ in transactions, demonstrating technical depth in complex systems",
        "Trained and mentored 7 engineers on technical fundamentals and customer communication best practices"
    ]

    for bullet in job3_bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)

    # Technical Skills section
    skills_heading = doc.add_paragraph()
    skills_run = skills_heading.add_run("TECHNICAL SKILLS")
    skills_run.font.bold = True
    skills_run.font.size = Pt(11)

    skills_text = doc.add_paragraph(
        "REST APIs • API Integration • JSON • Python • SQL • Postman • Browser Developer Tools • CloudWatch • Grafana • JIRA • Confluence • E-commerce Platforms • Product Data Configuration • API Documentation • SaaS Platforms"
    )
    skills_text.paragraph_format.space_after = Pt(12)

    # Languages section
    lang_heading = doc.add_paragraph()
    lang_run = lang_heading.add_run("LANGUAGES")
    lang_run.font.bold = True
    lang_run.font.size = Pt(11)

    lang_text = doc.add_paragraph(
        "English (C1/C2 – Professional Fluency) | Polish (Native)"
    )

    # Save the document
    resume_path = "/Users/bill/Work/codeworkspace/.claude/worktrees/interesting-swanson/job_agent/data/resumes/cv_linkedin_const_001.docx"
    doc.save(resume_path)
    print(f"Resume created successfully: {resume_path}")


def create_cover_letter():
    """Create the cover letter DOCX file"""
    doc = Document()

    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run("April 14, 2026")
    date_para.paragraph_format.space_after = Pt(12)

    # Recipient info (optional placeholder)
    recipient = doc.add_paragraph("Hiring Team\nConstructor\n")
    recipient.paragraph_format.space_after = Pt(12)

    # Salutation
    salutation = doc.add_paragraph("Dear Hiring Manager,")
    salutation.paragraph_format.space_after = Pt(12)

    # Opening paragraph
    opening = doc.add_paragraph(
        "I'm writing to express my interest in the Technical Integration Specialist role at Constructor. "
        "I'm currently at Luigi's Box working on AI-powered ecommerce search—exactly the problem Constructor solves. "
        "This domain expertise would enable me to onboard your customers with immediate product knowledge and confidence, making me effective from day one."
    )
    opening.paragraph_format.space_after = Pt(12)
    opening.paragraph_format.line_spacing = 1.15

    # Body paragraph 1
    body1 = doc.add_paragraph(
        "My integration and enterprise onboarding track record is exceptionally strong. At Luigi's Box, I've achieved zero churn rate "
        "from API integration implementations—a metric that reflects obsessive attention to customer success in every technical interaction. "
        "At Flexiroam, I led a successful B2B platform migration affecting 50+ enterprise accounts with a 100% implementation success rate. "
        "This isn't just support; it's technical mastery applied to complex, high-stakes integrations."
    )
    body1.paragraph_format.space_after = Pt(12)
    body1.paragraph_format.line_spacing = 1.15

    # Body paragraph 2
    body2 = doc.add_paragraph(
        "Throughout my 10+ years in SaaS, I've learned that the best integration specialist combines technical depth with customer advocacy. "
        "I don't just troubleshoot API issues—I understand the business context driving integration needs and ensure clients extract real value "
        "from their implementation. This approach has driven 25% year-over-year account growth and consistently delivered outcomes that exceed typical specialist-level expectations."
    )
    body2.paragraph_format.space_after = Pt(12)
    body2.paragraph_format.line_spacing = 1.15

    # Closing paragraph
    closing = doc.add_paragraph(
        "I'm drawn to Constructor because your growth-stage position in the ecommerce search market aligns perfectly with my domain expertise "
        "and proven ability to accelerate customer success. I'd welcome the opportunity to discuss how my integration mastery, API documentation expertise, "
        "and zero-churn track record can help Constructor build lasting customer relationships and scale your onboarding operation. "
        "Thank you for considering my application."
    )
    closing.paragraph_format.space_after = Pt(12)
    closing.paragraph_format.line_spacing = 1.15

    # Sign-off
    sign_off = doc.add_paragraph("Best regards,")
    sign_off.paragraph_format.space_after = Pt(24)

    # Name
    signature = doc.add_paragraph("Piotr Kaminski")
    signature.paragraph_format.space_after = Pt(0)

    # Save the document
    letter_path = "/Users/bill/Work/codeworkspace/.claude/worktrees/interesting-swanson/job_agent/data/cover_letters/letter_linkedin_const_001.docx"
    doc.save(letter_path)
    print(f"Cover letter created successfully: {letter_path}")


if __name__ == "__main__":
    try:
        create_resume()
        create_cover_letter()
        print("\nAll documents created successfully!")
    except Exception as e:
        print(f"Error creating documents: {e}")
        raise
