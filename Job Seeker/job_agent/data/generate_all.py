#!/usr/bin/env python3
import sys
sys.path.insert(0, '/Users/bill/Work/codeworkspace/.claude/worktrees/interesting-swanson/job_agent/data')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# ============================================================================
# CREATE RESUME
# ============================================================================
doc_resume = Document()

for section in doc_resume.sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

name = doc_resume.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_run = name.add_run('Piotr Kaminski')
name_run.font.size = Pt(14)
name_run.font.bold = True

contact = doc_resume.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.add_run('+48 510 934 702 | piotrkaminski1@protonmail.com | linkedin.com/in/piotr-kaminski | Krakow, Poland').font.size = Pt(9)

summary_h = doc_resume.add_paragraph()
summary_h.add_run('PROFESSIONAL SUMMARY').font.bold = True
summary_h.runs[0].font.size = Pt(11)

summary = doc_resume.add_paragraph('Technical Solutions Engineer with 10+ years driving API integration, SaaS product adoption, and customer success. Proven ability to troubleshoot complex technical issues, design customer onboarding processes, and deliver measurable business impact through cross-functional collaboration and data-driven problem solving.')
summary.paragraph_format.space_after = Pt(6)

exp_h = doc_resume.add_paragraph()
exp_h.add_run('EXPERIENCE').font.bold = True
exp_h.runs[0].font.size = Pt(11)

j1_title = doc_resume.add_paragraph()
j1_title.add_run('Lead Solutions Engineer').font.bold = True
j1_title.add_run(' | Luigi\'s Box | Jan 2025–Present')

for bullet in [
    'Achieved zero churn rate from integration implementations through comprehensive customer enablement and technical onboarding',
    'Led customer onboarding and integration initiatives for key accounts, ensuring 100% successful implementations',
    'Redesigned technical onboarding processes to improve customer success rates and reduce implementation friction',
    'Managed cross-functional collaboration between engineering and customer success teams; developed technical documentation for product enablement',
    'Conducted SQL data analysis to identify customer usage patterns and upsell opportunities'
]:
    p = doc_resume.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

j2_title = doc_resume.add_paragraph()
j2_title.add_run('Technical Account Manager').font.bold = True
j2_title.add_run(' | Flexiroam | Jan 2023–Dec 2024')

for bullet in [
    'Delivered 100% implementation success rate across all B2B migrations and customer integrations',
    'Grew assigned accounts by 25% year-over-year through technical enablement and customer support excellence',
    'Developed and maintained REST API integrations for enterprise clients; provided API troubleshooting and debugging support',
    'Led successful B2B platform migration affecting 50+ accounts; coordinated cross-functional team alignment',
    'Leveraged SQL and data analysis to drive customer adoption insights and upsell strategy'
]:
    p = doc_resume.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

j3_title = doc_resume.add_paragraph()
j3_title.add_run('Technical Support Engineer').font.bold = True
j3_title.add_run(' | Rapid | Jan 2020–Dec 2022')

for bullet in [
    'Managed Tier 2 technical support with 89% first-contact satisfaction rate; expertise in API troubleshooting and debugging complex integrations',
    'Developed and documented technical solutions for developer platform integration issues; created support documentation and runbooks',
    'Trained and mentored 7 engineers on API integration troubleshooting, technical support best practices, and customer communication',
    'Optimized support workflows and debugging processes, improving operational efficiency by 10-15% and reducing mean-time-to-resolution by 20%',
    'Developed fraud detection mechanisms protecting $300K+ in platform transactions through SQL query analysis'
]:
    p = doc_resume.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

skills_h = doc_resume.add_paragraph()
skills_h.add_run('TECHNICAL SKILLS & TOOLS').font.bold = True
skills_h.runs[0].font.size = Pt(11)

sk1 = doc_resume.add_paragraph()
sk1.add_run('APIs & Integration: ').font.bold = True
sk1.add_run('REST APIs, API troubleshooting, API integration, Webhooks, Integration support')

sk2 = doc_resume.add_paragraph()
sk2.add_run('Tools & Platforms: ').font.bold = True
sk2.add_run('Postman, SQL, JSON, CloudWatch, Grafana, JIRA, Confluence, VS Code')

sk3 = doc_resume.add_paragraph()
sk3.add_run('Languages: ').font.bold = True
sk3.add_run('English (C1/C2 professional fluency), Polish (native)')

sk4 = doc_resume.add_paragraph()
sk4.add_run('Core Competencies: ').font.bold = True
sk4.add_run('Technical troubleshooting, Customer onboarding, SaaS product knowledge, Cross-functional collaboration, Technical documentation, Problem solving, Process optimization, Team mentoring')

doc_resume.save('resumes/cv_linkedin_blz_001.docx')
print("✓ Resume created: resumes/cv_linkedin_blz_001.docx")

# ============================================================================
# CREATE COVER LETTER
# ============================================================================
doc_letter = Document()

for section in doc_letter.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

header = doc_letter.add_paragraph('Piotr Kaminski')
header.runs[0].font.bold = True
header.runs[0].font.size = Pt(12)

contact = doc_letter.add_paragraph('piotrkaminski1@protonmail.com | +48 510 934 702 | linkedin.com/in/piotr-kaminski')
contact.paragraph_format.space_after = Pt(8)

date_para = doc_letter.add_paragraph(datetime.now().strftime('%B %d, %Y'))
date_para.paragraph_format.space_after = Pt(12)

salutation = doc_letter.add_paragraph('Dear Hiring Team,')
salutation.paragraph_format.space_after = Pt(10)

opening = doc_letter.add_paragraph(
    'I\'m applying for the Technical Solutions Engineer role at Blazity because your focus on the AI/ML developer ecosystem '
    'directly aligns with my 10+ years building and scaling technical support for SaaS platforms. I\'ve spent the last decade mastering '
    'API integration troubleshooting, customer onboarding, and cross-functional collaboration—exactly what this role demands—and I\'m energized by the opportunity '
    'to bring that expertise to a rapidly growing AI/ML platform with a strong technical customer base.'
)
opening.paragraph_format.space_after = Pt(10)

middle1 = doc_letter.add_paragraph(
    'At RapidAPI (2020–2022), I managed Tier 2 technical support at scale, specializing in API troubleshooting and debugging complex integrations '
    'for a developer-first platform. I achieved an 89% first-contact satisfaction rate while handling escalations from across the platform. '
    'I trained 7 engineers on integration best practices and developed fraud detection mechanisms that protected $300K+ in transaction volume. '
    'That experience—supporting technical developers troubleshooting their own integrations—is directly transferable to Blazity\'s customer base.'
)
middle1.paragraph_format.space_after = Pt(10)

middle2 = doc_letter.add_paragraph(
    'As Technical Account Manager at Flexiroam, I grew assigned accounts by 25% YoY and delivered 100% implementation success across B2B migrations affecting 50+ accounts. '
    'I led the technical onboarding and API integration strategy for enterprise customers, coordinating closely with engineering and product teams. '
    'Most recently at Luigi\'s Box, I redesigned our customer onboarding processes and achieved zero churn on integration implementations. '
    'This trajectory—from support engineer to customer success to solutions engineer—demonstrates my ability to design and execute scalable customer enablement strategies.'
)
middle2.paragraph_format.space_after = Pt(10)

closing = doc_letter.add_paragraph(
    'I\'m based in Krakow and thrive in async-first, high-trust remote environments. I\'m excited to discuss how my API integration expertise, '
    'SQL-driven insights, and proven track record in customer enablement can help Blazity scale technical customer success. '
    'I\'d welcome the chance to explore this opportunity further.'
)
closing.paragraph_format.space_after = Pt(10)

sign_off = doc_letter.add_paragraph('Best regards,')
sign_off.paragraph_format.space_after = Pt(16)

name_sig = doc_letter.add_paragraph('Piotr Kaminski')

doc_letter.save('cover_letters/letter_linkedin_blz_001.docx')
print("✓ Cover letter created: cover_letters/letter_linkedin_blz_001.docx")

print("\n" + "="*70)
print("APPLICATION MATERIALS GENERATED SUCCESSFULLY")
print("="*70)
