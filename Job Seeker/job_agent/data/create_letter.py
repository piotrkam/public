#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Header
header = doc.add_paragraph('Piotr Kaminski')
header.runs[0].font.bold = True
header.runs[0].font.size = Pt(12)

contact = doc.add_paragraph('piotrkaminski1@protonmail.com | +48 510 934 702 | linkedin.com/in/piotr-kaminski')
contact.paragraph_format.space_after = Pt(8)

# Date
date_para = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
date_para.paragraph_format.space_after = Pt(12)

# Salutation
salutation = doc.add_paragraph('Dear Hiring Team,')
salutation.paragraph_format.space_after = Pt(10)

# Opening paragraph
opening = doc.add_paragraph(
    'I\'m applying for the Technical Solutions Engineer role at Blazity because your focus on the AI/ML developer ecosystem '
    'directly aligns with my 10+ years building and scaling technical support for SaaS platforms. I\'ve spent the last decade mastering '
    'API integration troubleshooting, customer onboarding, and cross-functional collaboration—exactly what this role demands—and I\'m energized by the opportunity '
    'to bring that expertise to a rapidly growing AI/ML platform with a strong technical customer base.'
)
opening.paragraph_format.space_after = Pt(10)

# Achievement 1
middle1 = doc.add_paragraph(
    'At RapidAPI (2020–2022), I managed Tier 2 technical support at scale, specializing in API troubleshooting and debugging complex integrations '
    'for a developer-first platform. I achieved an 89% first-contact satisfaction rate while handling escalations from across the platform. '
    'I trained 7 engineers on integration best practices and developed fraud detection mechanisms that protected $300K+ in transaction volume. '
    'That experience—supporting technical developers troubleshooting their own integrations—is directly transferable to Blazity\'s customer base.'
)
middle1.paragraph_format.space_after = Pt(10)

# Achievement 2
middle2 = doc.add_paragraph(
    'As Technical Account Manager at Flexiroam, I grew assigned accounts by 25% YoY and delivered 100% implementation success across B2B migrations affecting 50+ accounts. '
    'I led the technical onboarding and API integration strategy for enterprise customers, coordinating closely with engineering and product teams. '
    'Most recently at Luigi\'s Box, I redesigned our customer onboarding processes and achieved zero churn on integration implementations. '
    'This trajectory—from support engineer to customer success to solutions engineer—demonstrates my ability to design and execute scalable customer enablement strategies.'
)
middle2.paragraph_format.space_after = Pt(10)

# Closing
closing = doc.add_paragraph(
    'I\'m based in Krakow and thrive in async-first, high-trust remote environments. I\'m excited to discuss how my API integration expertise, '
    'SQL-driven insights, and proven track record in customer enablement can help Blazity scale technical customer success. '
    'I\'d welcome the chance to explore this opportunity further.'
)
closing.paragraph_format.space_after = Pt(10)

# Sign-off
sign_off = doc.add_paragraph('Best regards,')
sign_off.paragraph_format.space_after = Pt(16)

name_sig = doc.add_paragraph('Piotr Kaminski')

doc.save('cover_letters/letter_linkedin_blz_001.docx')
print("✓ Cover letter created")
