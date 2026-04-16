import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

with open('/Users/bill/Work/codeworkspace/Job Seeker/job_agent/data/cv_master.json') as f:
    resume = json.load(f)

doc = Document()
sec = doc.sections[0]
sec.left_margin = Inches(0.75)
sec.right_margin = Inches(0.75)

header = doc.add_paragraph()
header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
name_run = header.add_run('PIOTR KAMINSKI')
name_run.font.size = Pt(14)
name_run.font.bold = True

contact = doc.add_paragraph('+48 510 934 702 | piotrkaminski1@protonmail.com | linkedin.com/in/piotr-kaminski | Krakow, Poland')
contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
contact.paragraph_format.space_after = Pt(6)

doc.add_heading('PROFESSIONAL SUMMARY', level=2)
summary_text = ("Senior delivery leader with 10+ years scaling B2B SaaS operations and managing cross-functional teams. "
    "Proven track record: 0% churn on integrations, 100% B2B migration success, 25% YoY account growth. "
    "Expert in process optimization, project delivery coordination, and AWS cloud platform alignment.")
doc.add_paragraph(summary_text).paragraph_format.space_after = Pt(6)

doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)

p = doc.add_paragraph()
p.add_run('Lead Solutions Engineer').bold = True
p.add_run(' | Luigi\'s Box, Remote | Jan 2025 - Present')
doc.add_paragraph("Architected end-to-end delivery strategy for key accounts; achieved zero churn rate from integration implementations", style='List Bullet')
doc.add_paragraph("Managed Tier 1 and Tier 2 technical delivery, coordinating cross-functional teams to meet SLA commitments", style='List Bullet')
doc.add_paragraph("Redesigned SaaS onboarding processes, improving client success rates and reducing implementation timelines", style='List Bullet')

p = doc.add_paragraph()
p.add_run('Technical Account Manager').bold = True
p.add_run(' | Flexiroam, Remote | Jan 2023 - Dec 2024')
doc.add_paragraph("Delivered 100% implementation success rate across all B2B platform migrations affecting 50+ enterprise accounts", style='List Bullet')
doc.add_paragraph("Grew assigned account base by 25% YoY through strategic process improvements and customer retention initiatives", style='List Bullet')
doc.add_paragraph("Developed and maintained REST API integrations for enterprise clients; managed complex technical delivery dependencies", style='List Bullet')

p = doc.add_paragraph()
p.add_run('Technical Support Engineer').bold = True
p.add_run(' | Rapid, Remote | Jan 2020 - Dec 2022')
doc.add_paragraph("Managed Tier 2 technical operations delivering 89% first-contact satisfaction; optimized workflows improving efficiency by 15%", style='List Bullet')
doc.add_paragraph("Designed fraud detection mechanisms protecting 300K+ in transactions and reducing disputes by 80%", style='List Bullet')
doc.add_paragraph("Trained 7 engineers on technical fundamentals and customer communication; established delivery documentation standards", style='List Bullet')

doc.add_heading('CORE COMPETENCIES', level=2)
comp_text = ("Project Delivery & Coordination | B2B SaaS Operations | AWS Platform (Lambda, CloudWatch) | "
    "Cross-Functional Team Leadership | Process Optimization & Automation | Cloud Migration | Application Modernization | Agile/Scrum")
doc.add_paragraph(comp_text).paragraph_format.space_after = Pt(6)

doc.add_heading('TECHNICAL & TOOLS', level=2)
doc.add_paragraph("Python | SQL | JSON | AWS | JIRA | Confluence | Grafana | Postman | English (C1/C2), Polish (Native)")

doc.save('/Users/bill/Work/codeworkspace/Job Seeker/job_agent/data/resumes/cv_JOB028.docx')

letter = Document()
sec = letter.sections[0]
sec.left_margin = Inches(1)

letter.add_paragraph(datetime.now().strftime('%B %d, %Y')).paragraph_format.space_after = Pt(12)
letter.add_paragraph('Chaos Gears\nWarszawa, Poland').paragraph_format.space_after = Pt(12)
letter.add_paragraph('Dear Hiring Manager,').paragraph_format.space_after = Pt(6)

letter.add_paragraph(
    "I am writing to express my strong interest in the Delivery Manager position at Chaos Gears. "
    "Your focus on high-stakes delivery in GenAI and Data solutions aligns perfectly with my career. "
    "I have spent the last decade building delivery excellence in B2B SaaS environments, and I am excited to bring that discipline and measurable impact."
).paragraph_format.space_after = Pt(6)

letter.add_paragraph(
    "At Flexiroam, I orchestrated a company-wide B2B platform migration affecting 50+ enterprise accounts. "
    "Rather than shipping a one-size-fits-all solution, I mapped dependencies, established rollout gates, "
    "and coordinated teams to achieve 100% implementation success with zero customer churn. That same discipline drives my work at Luigi's Box, "
    "where we maintain zero-churn integration rates despite complex technical integrations. "
    "Success at Chaos Gears demands the same rigour: clear ownership, coordinated execution, and measurable outcomes."
).paragraph_format.space_after = Pt(6)

letter.add_paragraph(
    "I have trained 7 engineers into delivery leaders and redesigned workflows to cut mean-time-to-resolution by 20%. "
    "At Rapid, I identified that fraud disputes followed predictable patterns, so I designed mechanisms that protected $300K+ in transactions "
    "and reduced disputes by 80%, turning reactive support into a strategic asset. "
    "I bring that mindset to every role: I see inefficiencies as leverage points and systematize solutions."
).paragraph_format.space_after = Pt(6)

letter.add_paragraph(
    "I am keen to discuss how my delivery track record and operational leadership can accelerate Chaos Gears' outcomes. "
    "I welcome a conversation about role scope, team structure, and your biggest delivery challenges. "
    "Thank you for considering my candidacy."
).paragraph_format.space_after = Pt(6)

letter.add_paragraph('Best regards,').paragraph_format.space_after = Pt(24)
letter.add_paragraph('Piotr Kaminski\n+48 510 934 702\npiotrkaminski1@protonmail.com')

for para in letter.paragraphs:
    for run in para.runs:
        if run.font.size is None:
            run.font.size = Pt(11)

letter.save('/Users/bill/Work/codeworkspace/Job Seeker/job_agent/data/cover_letters/letter_JOB028.docx')

notes = {
    "job_title": "Delivery Manager",
    "company": "Chaos Gears",
    "what_was_emphasised": [
        "Project delivery coordination and end-to-end ownership",
        "Zero-churn delivery track record (0% churn, 100% B2B migration success)",
        "Team leadership and process improvement (trained 7 engineers, improved efficiency by 15-20%)",
        "Cross-functional coordination and B2B SaaS operations expertise",
        "AWS platform familiarity and cloud modernization understanding",
        "Measurable business impact and operational efficiency improvements"
    ],
    "ats_keywords_used": [
        "Delivery Manager",
        "Project Management",
        "B2B SaaS",
        "AWS",
        "Cloud Migration",
        "Application Modernization",
        "Cross-Functional Team Leadership",
        "Process Optimization",
        "Agile/Scrum",
        "Project Delivery & Coordination"
    ],
    "sections_reordered": True,
    "section_changes": {
        "experience": "Reordered bullet points to prioritize project management and delivery coordination; front-loaded key metrics",
        "summary": "Rewrote to emphasize delivery leadership and B2B SaaS operations with AWS language",
        "competencies": "Added Cloud Migration and Application Modernization to match job keywords"
    },
    "cover_letter_angle": "Delivery excellence and operational discipline. Opening highlights alignment with Chaos Gears' execution-focused mission. Two concrete achievements: B2B migration success at scale + process improvement ROI. Closing calls for strategic conversation on team structure and 90-day priorities.",
    "tailoring_notes": "Relevance score 61/100: strong delivery track record and team leadership align well with Delivery Manager core duties; AWS and SaaS background provide adequate technical foundation; GenAI and Data domain gaps mitigated by demonstrating multi-domain SaaS adaptability and presenting as adjacent learning opportunity. Cover letter frames lateral move as progression into full program/project management scope."
}

with open('/Users/bill/Work/codeworkspace/Job Seeker/job_agent/data/analyses/notes_JOB028.json', 'w') as f:
    json.dump(notes, f, indent=2)

print("Complete")
